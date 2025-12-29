import streamlit as st
import google.generativeai as genai
from docx import Document
from streamlit_mermaid import st_mermaid
from audio_recorder_streamlit import audio_recorder
import tempfile
import os
import time
import mimetypes
import re
import random

# --- 1. CẤU HÌNH TRANG ---
st.set_page_config(page_title="Universal AI Studio (Ultimate)", page_icon="💎", layout="wide")
st.markdown("""
<style>
    .stButton>button {width: 100%; border-radius: 8px; height: 3em; font-weight: bold; background: #1e3c72; color: white;}
    .stExpander {border: 1px solid #e0e0e0; border-radius: 8px; margin-bottom: 10px; background-color: #ffffff;}
    .stMarkdown h2 {color: #1a2a6c; border-bottom: 2px solid #eee; padding-bottom: 5px;}
    .error-box {padding: 15px; background-color: #ffebee; border: 1px solid #ffcdd2; border-radius: 5px; color: #c62828; margin-bottom: 10px;}
</style>
""", unsafe_allow_html=True)

# --- 2. BIẾN TOÀN CỤC ---
STRICT_RULES = "CHỈ DÙNG FILE GỐC. CẤM BỊA TÊN DIỄN GIẢ. CẤM BỊA NỘI DUNG. TRÍCH DẪN GIỜ [mm:ss]."

# --- 3. QUẢN LÝ SESSION ---
if "chat_history" not in st.session_state: st.session_state.chat_history = []
if "gemini_files" not in st.session_state: st.session_state.gemini_files = [] 
if "analysis_result" not in st.session_state: st.session_state.analysis_result = ""
if "is_auto_running" not in st.session_state: st.session_state.is_auto_running = False
if "loop_count" not in st.session_state: st.session_state.loop_count = 0
# Biến phục vụ Retry
if "quota_error" not in st.session_state: st.session_state.quota_error = False
if "last_prompt" not in st.session_state: st.session_state.last_prompt = ""
if "last_config" not in st.session_state: st.session_state.last_config = None

# --- 4. HÀM HỖ TRỢ ---
def get_system_key():
    try:
        if "SYSTEM_KEYS" in st.secrets:
            keys = st.secrets["SYSTEM_KEYS"]
            if isinstance(keys, str): 
                keys = [k.strip() for k in keys.replace('[','').replace(']','').replace('"','').replace("'",'').split(',')]
            return random.choice(keys)
        elif "GOOGLE_API_KEY" in st.secrets:
            return st.secrets["GOOGLE_API_KEY"]
    except: return None

def configure_genai(user_key=None):
    api_key = user_key if user_key else get_system_key()
    if not api_key: return False
    try:
        genai.configure(api_key=api_key)
        return True
    except: return False

def get_optimized_models():
    # Danh sách cứng để đảm bảo luôn có lựa chọn
    return ["models/gemini-3.0-flash-preview", "models/gemini-2.0-flash-exp", "models/gemini-1.5-flash", "models/gemini-1.5-pro"]

def format_model_name(name):
    return name.replace("models/", "").replace("-preview", " (Pre)").replace("-latest", "").upper()

def upload_to_gemini(path):
    mime_type, _ = mimetypes.guess_type(path)
    file = genai.upload_file(path, mime_type=mime_type or "application/octet-stream")
    while file.state.name == "PROCESSING":
        time.sleep(1)
        file = genai.get_file(file.name)
    return file

def create_docx(content):
    doc = Document()
    doc.add_heading('BÁO CÁO', 0)
    clean_content = content.replace("```markdown", "").replace("```", "")
    for line in clean_content.split('\n'):
        if line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
        else: doc.add_paragraph(line)
    return doc

def get_safe_response(response):
    try:
        finish_reason = response.candidates[0].finish_reason
        if finish_reason in [1, 2]: return response.text
        elif finish_reason == 3: return "\n\n[CẢNH BÁO: Nội dung bị chặn do Safety.]"
        elif finish_reason == 4: return "\n\n[DỪNG: Phát hiện nội dung có bản quyền.]"
        else: return f"\n\n[Lỗi: Finish Reason {finish_reason}]"
    except: return response.text

# --- 5. MAIN APP ---
def main():
    st.title("💎 Universal AI Studio (Ultimate)")
    
    # --- SIDEBAR ---
    with st.sidebar:
        st.header("🎯 CHẾ ĐỘ")
        main_mode = st.radio("Mục tiêu:", ("📝 Gỡ băng nguyên văn", "📊 Phân tích chuyên sâu"))
        
        if main_mode == "📊 Phân tích chuyên sâu":
            st.subheader("KHO VŨ KHÍ (FULL):")
            
            st.markdown("**1. Cốt lõi**")
            opt_summary = st.checkbox("📋 Tóm tắt & Hành động", True)
            opt_process = st.checkbox("🔄 Trích xuất Quy trình", False)
            opt_prosody = st.checkbox("🎭 Phân tích Cảm xúc", False)
            opt_gossip = st.checkbox("☕ Chế độ 'Bà tám'", False)

            st.markdown("**2. Sáng tạo Nghe/Nhìn**")
            opt_podcast = st.checkbox("🎙️ Kịch bản Podcast", False)
            opt_video = st.checkbox("🎬 Kịch bản Video", False)
            opt_mindmap = st.checkbox("🧠 Sơ đồ tư duy (Mindmap)", True)

            st.markdown("**3. Học tập & Nghiên cứu**")
            opt_report = st.checkbox("📑 Báo cáo chuyên sâu", False)
            opt_briefing = st.checkbox("📄 Briefing Doc (Tóm lược)", False)
            opt_timeline = st.checkbox("⏳ Timeline (Dòng thời gian)", False)
            opt_faq = st.checkbox("❓ FAQ (Hỏi đáp)", False)
            opt_quiz = st.checkbox("📝 Quiz & Thẻ nhớ", False)
            
            st.markdown("**4. Dữ liệu**")
            opt_slides = st.checkbox("🖥️ Dàn ý Slide", False)
            opt_table = st.checkbox("📉 Bảng số liệu", False)
        
        st.divider()
        
        with st.expander("⚙️ Cấu hình & Key", expanded=True):
            initial_key = st.text_input("Key riêng (Tùy chọn):", type="password")
            if configure_genai(initial_key):
                st.success("Đã kết nối!")
                models = get_optimized_models()
                model_version = st.selectbox("Engine:", models, index=0, format_func=format_model_name)
                if main_mode.startswith("📊"):
                    detail_level = st.select_slider("Chi tiết:", ["Sơ lược", "Tiêu chuẩn", "Sâu"], value="Sâu")
            else: st.error("Chưa kết nối!")

        if st.button("🗑️ Reset"):
            st.session_state.clear(); st.rerun()

    # --- XỬ LÝ LỖI QUOTA (INTERACTIVE) ---
    if st.session_state.quota_error:
        st.markdown("""
        <div class="error-box">
            <h3>⚠️ HẾT HẠN MỨC (429 QUOTA EXCEEDED)</h3>
            <p>Model bạn chọn đã hết lượt dùng miễn phí. Bạn muốn xử lý sao?</p>
        </div>
        """, unsafe_allow_html=True)
        
        c1, c2 = st.columns(2)
        with c1:
            rescue_key = st.text_input("🔑 Nhập API Key MỚI để tiếp tục:", type="password", key="rescue")
            if st.button("🚀 Thử lại với Key này"):
                if configure_genai(rescue_key):
                    st.session_state.quota_error = False
                    st.rerun() # Chạy lại với key mới
        with c2:
            st.write("Hoặc:")
            if st.button("⬇️ Hạ xuống 1.5 Flash (Miễn phí)"):
                st.session_state.quota_error = False
                with st.spinner("Đang chuyển sang 1.5 Flash..."):
                    try:
                        # Hạ cấp model
                        model = genai.GenerativeModel("models/gemini-1.5-flash")
                        response = model.generate_content([st.session_state.last_prompt] + st.session_state.gemini_files, generation_config=st.session_state.last_config)
                        st.session_state.analysis_result = get_safe_response(response)
                        st.rerun()
                    except Exception as e: st.error(f"Lỗi: {e}")
        st.divider()

    # --- TABS ---
    tab_work, tab_chat = st.tabs(["📂 Xử lý", "💬 Chat"])

    with tab_work:
        if not st.session_state.quota_error:
            up_files = st.file_uploader("Upload file", accept_multiple_files=True)
            audio_bytes = audio_recorder()

            if st.button("🚀 BẮT ĐẦU", type="primary"):
                temp_paths = []
                if up_files:
                    for f in up_files:
                        ext = os.path.splitext(f.name)[1] or ".txt"
                        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                            tmp.write(f.getvalue()); temp_paths.append(tmp.name)
                if audio_bytes:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
                        tmp.write(audio_bytes); temp_paths.append(tmp.name)
                
                if not temp_paths:
                    st.warning("Chưa có file!")
                else:
                    with st.spinner(f"Đang xử lý với {format_model_name(model_version)}..."):
                        try:
                            g_files = [upload_to_gemini(p) for p in temp_paths]
                            st.session_state.gemini_files = g_files
                            
                            # Tắt bộ lọc an toàn
                            safety_settings = [
                                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
                            ]
                            
                            gen_config = genai.types.GenerationConfig(max_output_tokens=8192, temperature=0.2)

                            if main_mode.startswith("📝"):
                                prompt = f"""
                                {STRICT_RULES}
                                NHIỆM VỤ: Gỡ băng NGUYÊN VĂN 100%.
                                YÊU CẦU:
                                1. Bắt đầu mỗi câu bằng [Phút:Giây].
                                2. Viết lại chính xác từng từ.
                                3. Định danh: 'Người nói 1', 'Người nói 2'.
                                4. Ngôn ngữ: Tiếng Việt.
                                """
                                # MẶC ĐỊNH BẬT AUTO CONTINUE
                                st.session_state.is_auto_running = True
                                st.session_state.loop_count = 1
                            else:
                                # TỔNG HỢP PROMPT FULL TÍNH NĂNG
                                prompt = f"{STRICT_RULES}\nNHIỆM VỤ: Phân tích sâu {detail_level} cho các mục sau:\n"
                                if opt_summary: prompt += "## TÓM TẮT & HÀNH ĐỘNG\n"
                                if opt_process: prompt += "## QUY TRÌNH CHI TIẾT\n"
                                if opt_prosody: prompt += "## PHÂN TÍCH CẢM XÚC\n"
                                if opt_gossip: prompt += "## GÓC BÀ TÁM\n"
                                if opt_podcast: prompt += "## KỊCH BẢN PODCAST\n"
                                if opt_video: prompt += "## KỊCH BẢN VIDEO\n"
                                if opt_mindmap: prompt += "## MÃ SƠ ĐỒ TƯ DUY (Mermaid)\n"
                                if opt_report: prompt += "## BÁO CÁO CHUYÊN SÂU\n"
                                if opt_briefing: prompt += "## BRIEFING DOC\n"
                                if opt_timeline: prompt += "## TIMELINE SỰ KIỆN\n"
                                if opt_faq: prompt += "## CÂU HỎI THƯỜNG GẶP (FAQ)\n"
                                if opt_quiz: prompt += "## TRẮC NGHIỆM & THẺ NHỚ\n"
                                if opt_slides: prompt += "## DÀN Ý SLIDE\n"
                                if opt_table: prompt += "## BẢNG SỐ LIỆU\n"

                            # LƯU TRẠNG THÁI ĐỂ RETRY
                            st.session_state.last_prompt = prompt
                            st.session_state.last_config = gen_config

                            model = genai.GenerativeModel(model_version)
                            response = model.generate_content(
                                [prompt] + g_files, 
                                generation_config=gen_config,
                                safety_settings=safety_settings
                            )
                            
                            st.session_state.analysis_result = get_safe_response(response)
                            st.rerun()
                            
                        except Exception as e:
                            err_msg = str(e)
                            if "429" in err_msg or "Quota" in err_msg:
                                st.session_state.quota_error = True
                                st.rerun()
                            elif "404" in err_msg or "Not Found" in err_msg:
                                st.toast(f"⚠️ Model {format_model_name(model_version)} lỗi (404). Tự động chuyển sang 1.5 Flash...", icon="🔄")
                                try:
                                    fb_model = genai.GenerativeModel("models/gemini-1.5-flash")
                                    response = fb_model.generate_content([prompt] + g_files, generation_config=gen_config, safety_settings=safety_settings)
                                    st.session_state.analysis_result = get_safe_response(response)
                                    st.rerun()
                                except Exception as e2: st.error(f"Lỗi hệ thống: {e2}")
                            else:
                                st.error(f"Lỗi: {e}")

        # HIỂN THỊ KẾT QUẢ
        if st.session_state.analysis_result:
            if st.session_state.is_auto_running:
                st.warning(f"🔄 Đang tự động chạy tiếp (Vòng {st.session_state.loop_count})...")
                if st.button("🛑 DỪNG"):
                    st.session_state.is_auto_running = False
                    st.success("Đã dừng."); st.rerun()

            st.divider()
            res = st.session_state.analysis_result
            
            # Xử lý Mindmap
            if "```mermaid" in res:
                try:
                    m_code = res.split("```mermaid")[1].split("```")[0]
                    st_mermaid(m_code, height=500)
                except: pass
            
            # Hiển thị Text
            sections = res.split("## ")
            for s in sections:
                if not s.strip(): continue
                lines = s.split("\n")
                with st.expander(f"📌 {lines[0].strip()}", expanded=True):
                    st.markdown("\n".join(lines[1:]))

            # Download
            doc = create_docx(res)
            doc_io = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(doc_io.name)
            with open(doc_io.name, "rb") as f:
                st.download_button("📥 Tải Báo Cáo", f, "Bao_Cao.docx", type="primary")
            os.remove(doc_io.name)

            # AUTO-CONTINUE (Logic cũ + Retry Quota)
            if st.session_state.is_auto_running and main_mode.startswith("📝"):
                if "[DỪNG:" in res or "[CẢNH BÁO:" in res:
                    st.session_state.is_auto_running = False
                    st.error("⚠️ Dừng do bản quyền/an toàn.")
                else:
                    st.divider()
                    placeholder = st.empty()
                    for i in range(3, 0, -1):
                        placeholder.info(f"⏳ Chạy tiếp trong {i}s...")
                        time.sleep(1)
                    placeholder.empty()
                    
                    with st.spinner("Đang nghe tiếp..."):
                        try:
                            cont_config = genai.types.GenerationConfig(max_output_tokens=8192, temperature=0.2)
                            model = genai.GenerativeModel(model_version)
                            last_part = res[-500:]
                            c_prompt = f"""
                            CONTEXT: Đang gỡ băng dở dang.
                            MỎ NEO: "...{last_part}"
                            NHIỆM VỤ: Tìm mỏ neo, viết tiếp NGUYÊN VĂN đoạn sau. KHÔNG viết lại mỏ neo.
                            """
                            
                            safety_settings = [
                                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
                            ]
                            
                            # Lưu trạng thái cho Retry
                            st.session_state.last_prompt = c_prompt
                            st.session_state.last_config = cont_config

                            c_res = model.generate_content(
                                [c_prompt] + st.session_state.gemini_files, 
                                generation_config=cont_config,
                                safety_settings=safety_settings
                            )
                            
                            safe_c_text = get_safe_response(c_res)

                            if len(safe_c_text) < 50 or "kết thúc" in safe_c_text.lower() or "[DỪNG:" in safe_c_text:
                                st.session_state.is_auto_running = False
                                st.success("✅ Đã xong!")
                                if "[DỪNG:" in safe_c_text:
                                    st.session_state.analysis_result += "\n\n" + safe_c_text
                                    st.rerun()
                            else:
                                st.session_state.analysis_result += "\n\n" + safe_c_text
                                st.session_state.loop_count += 1
                                st.rerun()
                        except Exception as e:
                            err_msg = str(e)
                            if "429" in err_msg or "Quota" in err_msg:
                                st.session_state.quota_error = True
                                st.session_state.is_auto_running = False # Dừng auto để xử lý lỗi
                                st.rerun()
                            elif "404" in err_msg:
                                # 404 thì tự fallback luôn
                                try:
                                    fb_model = genai.GenerativeModel("models/gemini-1.5-flash")
                                    c_res = fb_model.generate_content([c_prompt] + st.session_state.gemini_files, generation_config=cont_config, safety_settings=safety_settings)
                                    st.session_state.analysis_result += "\n\n" + get_safe_response(c_res)
                                    st.session_state.loop_count += 1
                                    st.rerun()
                                except: st.error("Lỗi hệ thống."); st.session_state.is_auto_running = False
                            else:
                                st.error(f"Lỗi: {e}")
                                st.session_state.is_auto_running = False

    with tab_chat:
        st.header("💬 Chat")
        if st.session_state.gemini_files:
            for m in st.session_state.chat_history:
                with st.chat_message(m["role"]): st.markdown(m["content"])
            if inp := st.chat_input("Hỏi AI..."):
                st.session_state.chat_history.append({"role": "user", "content": inp})
                with st.chat_message("user"): st.markdown(inp)
                with st.chat_message("assistant"):
                    try:
                        m = genai.GenerativeModel(model_version)
                        r = m.generate_content(
                            st.session_state.gemini_files + [f"Trả lời: {inp}"],
                            safety_settings=[{"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"}]
                        )
                        st.markdown(r.text); st.session_state.chat_history.append({"role": "assistant", "content": r.text})
                    except: st.error("Lỗi chat.")
        else: st.info("👈 Upload file trước.")

if __name__ == "__main__":
    main()
